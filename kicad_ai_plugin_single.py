import os
import wx  # type: ignore
import pcbnew  # type: ignore
import logging
import threading
import json
import time
import urllib.request
import urllib.error
import ssl
import http.client
import uuid
import datetime
import base64
from typing import Dict, List, Any, Optional

# Import KiCad file processing utilities
try:
    from kicad_file_processor import is_kicad_file, extract_kicad_file_info
    HAS_KICAD_PROCESSOR = True
except ImportError:
    HAS_KICAD_PROCESSOR = False
    # Simple fallback functions
    def is_kicad_file(file_ext):
        return file_ext.lower() in ['.kicad_pcb', '.kicad_sch', '.kicad_pro', '.net', 
                                   '.lib', '.kicad_mod', '.kicad_wks', '.kicad_sym']
    def extract_kicad_file_info(filepath, file_ext, content):
        return {"file_type": "KiCad File", "summary": "KiCad File (detailed info not available)"}

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# No default API key - user must provide their own
DEFAULT_API_KEY = ""
DEFAULT_MODEL = "gpt-3.5-turbo"  # Using a more cost-effective model as default

# Detailed system prompt for better AI assistant guidance
SYSTEM_PROMPT = """
You are an intelligent AI assistant integrated into a KiCad plugin called "KiCad AI Plugin by ALT TAB".

Your primary abilities include:
1. Answering questions on a wide range of topics with accurate information
2. Analyzing uploaded files and providing insights
3. Maintaining helpful, conversational interactions
4. Providing clear explanations of complex concepts
5. Remembering context throughout a conversation
6. Answer in the language of the inputted prompt or how the user specifies in the prompt, default is English.

When responding to questions about KiCad or electronics:
- Provide clear, concise explanations
- Include step-by-step instructions when appropriate
- Reference specific UI elements or menu paths when relevant
- Suggest best practices based on industry standards

For file analysis:
- Identify the file type and its purpose
- Highlight key elements or potential issues
- Explain the content in user-friendly terms

Maintain a helpful, friendly, and professional tone in all interactions.
"""

# Configuration file path for API key
API_CONFIG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config", "api_config.json")

# Maximum file size (5MB)
MAX_FILE_SIZE = 5 * 1024 * 1024

# Helper functions for API key management
def get_api_key():
    """Get API key from config file"""
    try:
        # Create the config directory if it doesn't exist
        config_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config")
        os.makedirs(config_dir, exist_ok=True)
        
        # Check if the config file exists
        if os.path.exists(API_CONFIG_FILE):
            with open(API_CONFIG_FILE, "r") as f:
                config = json.load(f)
                return config.get("api_key", "")
        return ""
    except Exception as e:
        logger.error(f"Error loading API key: {str(e)}")
        return ""

def get_base_url():
    """Get base URL from config file"""
    try:
        if os.path.exists(API_CONFIG_FILE):
            with open(API_CONFIG_FILE, "r") as f:
                config = json.load(f)
                return config.get("base_url", "api.openai.com")
        return "api.openai.com"
    except Exception as e:
        logger.error(f"Error loading base URL: {str(e)}")
        return "api.openai.com"

def save_api_key(api_key):
    """Save API key to config file"""
    try:
        # Create the config directory if it doesn't exist
        config_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config")
        os.makedirs(config_dir, exist_ok=True)
        
        # Keep existing config like base_url
        existing_config = {}
        if os.path.exists(API_CONFIG_FILE):
            try:
                with open(API_CONFIG_FILE, "r") as f:
                    existing_config = json.load(f)
            except Exception:
                pass
                
        existing_config["api_key"] = api_key
        
        # Save the config
        with open(API_CONFIG_FILE, "w") as f:
            json.dump(existing_config, f, indent=2)
        return True
    except Exception as e:
        logger.error(f"Error saving API key: {str(e)}")
        return False

class Conversation:
    """Class to store a single conversation"""
    def __init__(self, id=None, title="New Chat", messages=None):
        self.id = id or str(uuid.uuid4())
        self.title = title
        self.messages = messages or []
        self.created_at = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.last_updated = self.created_at
        self.files = {}  # Store file data separately to avoid bloating the messages
        
    def add_message(self, role, content, file_data=None):
        """Add a message to the conversation"""
        # If there's file data, store it separately with a reference ID
        file_ref = None
        if file_data:
            file_id = str(uuid.uuid4())
            self.files[file_id] = file_data
            file_ref = file_id
            
        # Add the message with optional file reference
        message = {"role": role, "content": content}
        if file_ref:
            message["file_ref"] = file_ref
            
        self.messages.append(message)
        self.last_updated = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Update the title based on the first user message (auto-naming)
        if role == "user" and len([m for m in self.messages if m["role"] == "user"]) == 1:
            # Only use first user message to set title (auto-naming)
            # Extract a clean title without file content
            clean_content = content
            logger.info(f"Auto-naming conversation based on first message: {clean_content[:30]}...")
            
            if "[Attached file:" in clean_content:
                # Get the text before the file attachment
                parts = clean_content.split("[Attached file:")
                clean_content = parts[0].strip()
                if not clean_content:
                    # Use the filename as the title if there's no message
                    filename_end = parts[1].find("]")
                    if filename_end > 0:
                        clean_content = "File: " + parts[1][:filename_end].strip()
            
            # Create a title that's not too long but still descriptive
            if clean_content:
                # Limit to the first sentence or 50 characters, whichever is shorter
                sentences = clean_content.split('.')
                first_sentence = sentences[0].strip()
                
                # Ensure it's not too long
                if len(first_sentence) > 50:
                    title = first_sentence[:47] + "..."
                else:
                    title = first_sentence
                    
                # Add ellipsis if needed
                if len(title) < len(clean_content) and not title.endswith("..."):
                    title += "..."
                    
                self.title = title
                logger.info(f"Updated conversation title to: '{self.title}'")
            else:
                logger.info("No content for auto-naming, keeping default title")
    
    def rename(self, new_title):
        """Rename the conversation"""
        self.title = new_title
        self.last_updated = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
    def to_dict(self):
        """Convert to dictionary for serialization"""
        return {
            "id": self.id,
            "title": self.title,
            "messages": self.messages,
            "created_at": self.created_at,
            "last_updated": self.last_updated,
            "files": self.files
        }
        
    @classmethod
    def from_dict(cls, data):
        """Create a conversation from a dictionary"""
        conv = cls(
            id=data.get("id"),
            title=data.get("title", "New Chat"),
            messages=data.get("messages", [])
        )
        conv.created_at = data.get("created_at", conv.created_at)
        conv.last_updated = data.get("last_updated", conv.last_updated)
        conv.files = data.get("files", {})
        return conv

class ConversationManager:
    """Manages multiple conversations"""
    def __init__(self):
        self.conversations = {}
        self.current_conversation_id = None
        self.load_conversations()
        
    def new_conversation(self, system_message=None):
        """Create a new conversation"""
        conversation = Conversation()
        if system_message:
            conversation.add_message("system", system_message)
        self.conversations[conversation.id] = conversation
        self.current_conversation_id = conversation.id
        self.save_conversations()
        return conversation
        
    def get_current_conversation(self):
        """Get the current conversation"""
        if not self.current_conversation_id or self.current_conversation_id not in self.conversations:
            return self.new_conversation()
        return self.conversations[self.current_conversation_id]
        
    def switch_conversation(self, conversation_id):
        """Switch to a different conversation"""
        if conversation_id in self.conversations:
            self.current_conversation_id = conversation_id
            return self.conversations[conversation_id]
        return None
        
    def delete_conversation(self, conversation_id):
        """Delete a conversation by ID."""
        try:
            if conversation_id not in self.conversations:
                logger.warning(f"Attempted to delete non-existent conversation: {conversation_id}")
                return False
                
            # If we're deleting the current conversation, switch to another one first
            if self.current_conversation_id == conversation_id:
                # Find an alternative conversation
                alternative_ids = [cid for cid in self.conversations.keys() if cid != conversation_id]
                
                if alternative_ids:
                    # Switch to another conversation before deleting
                    self.current_conversation_id = alternative_ids[0]
                    logger.info(f"Switched to conversation {self.current_conversation_id} before deleting {conversation_id}")
                else:
                    # Create a new conversation if this is the last one
                    new_conversation = self.new_conversation()
                    self.current_conversation_id = new_conversation.id
                    logger.info(f"Created new conversation {new_conversation.id} before deleting {conversation_id}")
            
            # Now it's safe to delete
            del self.conversations[conversation_id]
            logger.info(f"Successfully deleted conversation {conversation_id}")
            self.save_conversations()
            return True
        except Exception as e:
            logger.error(f"Error deleting conversation {conversation_id}: {str(e)}")
            # Ensure we have a valid current conversation ID
            if self.current_conversation_id not in self.conversations and self.conversations:
                # Set to any existing conversation
                self.current_conversation_id = list(self.conversations.keys())[0]
                logger.info(f"Set current conversation to {self.current_conversation_id} after deletion error")
            elif not self.conversations:
                # Create a new conversation if none exist
                new_conversation = self.new_conversation()
                self.current_conversation_id = new_conversation.id
                logger.info(f"Created new conversation {new_conversation.id} after deletion error recovery")
            return False
            
    def rename_conversation(self, conversation_id, new_title):
        """Rename a conversation"""
        try:
            if conversation_id in self.conversations:
                self.conversations[conversation_id].title = new_title
                self.save_conversations()
                return True
            return False
        except Exception as e:
            logger.error(f"Error renaming conversation: {str(e)}")
            return False
            
    def save_conversations(self):
        """Save conversations to disk"""
        try:
            data = {
                "current_id": self.current_conversation_id,
                "conversations": {id: conv.to_dict() for id, conv in self.conversations.items()}
            }
            
            # Create the conversations directory if it doesn't exist
            conversations_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "conversations")
            os.makedirs(conversations_dir, exist_ok=True)
            
            # Save the conversations
            path = os.path.join(conversations_dir, "conversations.json")
            with open(path, 'w') as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving conversations: {str(e)}")
            
    def load_conversations(self):
        """Load conversations from disk"""
        try:
            # Get the conversations path
            conversations_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "conversations")
            path = os.path.join(conversations_dir, "conversations.json")
            
            # Check if the file exists
            if not os.path.exists(path):
                return
                
            with open(path, 'r') as f:
                data = json.load(f)
                
            self.current_conversation_id = data.get("current_id")
            
            for conv_id, conv_data in data.get("conversations", {}).items():
                self.conversations[conv_id] = Conversation.from_dict(conv_data)
                
            # If no conversations were loaded, create a new one
            if not self.conversations:
                self.new_conversation()
                
        except Exception as e:
            logger.error(f"Error loading conversations: {str(e)}")
            # Create a new conversation if loading failed
            self.new_conversation()

class OpenAIClient:
    """
    Simple OpenAI client that doesn't require the openai package
    Using direct HTTP requests to the OpenAI API
    """
    
    def __init__(self, api_key=None):
        self.api_key = api_key or get_api_key()
        self.base_url = get_base_url()
        self.model = DEFAULT_MODEL
        
    def set_model(self, model):
        """Set the model to use"""
        self.model = model
        
    def chat_completion(self, messages, conversation_files=None):
        """
        Make a chat completion request to the OpenAI API
        Implements the API directly without requiring the openai package
        
        Args:
            messages: List of message objects
            conversation_files: Dictionary of file references from the conversation
        """
        try:
            # Check if API key is set
            if not self.api_key:
                return "Error: No API key provided. Please configure your OpenAI API key."
                
            # For models that support it, we can add image attachments
            supports_vision = self.model in ["gpt-4-vision-preview", "gpt-4-turbo", "gpt-4o", "gpt-4-1106-vision-preview"]
            
            # Process messages to convert binary files to base64 for vision-capable models
            processed_messages = []
            
            # Process each message
            for msg in messages:
                # Handle system messages
                if msg["role"] == "system":
                    processed_messages.append({"role": msg["role"], "content": msg["content"]})
                    continue
                
                # First add the message as-is
                processed_messages.append({"role": msg["role"], "content": msg["content"]})
                
                # Then check if we need special handling for images
                if not supports_vision:
                    continue
                
                # Check if this message has a file reference
                if "file_ref" not in msg:
                    continue
                    
                # Check if we have the files dictionary
                if not conversation_files:
                    continue
                    
                # Get the file reference
                file_ref = msg["file_ref"]
                if file_ref not in conversation_files:
                    continue
                    
                # Get the file data
                file_data = conversation_files[file_ref]
                
                # Check if this is an image file
                if not file_data.get('is_binary'):
                    continue
                    
                if not file_data.get('content_b64'):
                    continue
                    
                if not file_data.get('mime_type', '').startswith('image/'):
                    continue
                    
                # This is an image file, so handle it specially for vision models
                logger.info(f"Converting image file to vision format: {file_data.get('name')}")
                    
                # Get message content
                message_text = msg["content"]
                    
                # Create a content array for vision API
                content = [
                    {"type": "text", "text": message_text},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{file_data['mime_type']};base64,{file_data['content_b64']}",
                            "detail": "auto"
                        }
                    }
                ]
                    
                # Replace the last added message with the vision format
                processed_messages.pop()
                processed_messages.append({"role": msg["role"], "content": content})
            
            # Use processed messages for API request
            data = {
                "model": self.model,
                "messages": processed_messages,
                "temperature": 0.0  # Use deterministic responses for technical questions
            }
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            
            # Create an SSL context that doesn't verify certificates
            # This is not ideal for production but helps bypass SSL issues
            context = ssl._create_unverified_context()
            
            # Create a connection with the custom SSL context
            conn = http.client.HTTPSConnection(self.base_url, context=context)
            conn.request("POST", "/v1/chat/completions", json.dumps(data), headers)
            response = conn.getresponse()
            
            if response.status != 200:
                error_content = response.read().decode()
                logger.error(f"Error from OpenAI API: {error_content}")
                return f"Error from OpenAI API: {response.status} - {error_content}"
                
            response_data = json.loads(response.read().decode())
            return response_data["choices"][0]["message"]["content"]
            
        except Exception as e:
            logger.error(f"Error calling OpenAI API: {str(e)}")
            return f"Error calling OpenAI API: {str(e)}"

class OpenAIManager:
    """Manages OpenAI API interactions"""
    
    def __init__(self):
        # Load API key from config
        self.api_key = get_api_key()
        self.model = DEFAULT_MODEL
        self.system_prompt = SYSTEM_PROMPT
        self.initialized = True
        self.client = OpenAIClient(self.api_key)
        self.client.set_model(self.model)
        
    def load_config(self):
        """Load configuration from file"""
        config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config", "openai_config.json")
        try:
            if os.path.exists(config_path):
                with open(config_path, 'r') as f:
                    config = json.load(f)
                    self.model = config.get("model", DEFAULT_MODEL)
                    self.system_prompt = config.get("system_prompt", SYSTEM_PROMPT)
                    
            # Always load the API key from the API config file
            self.api_key = get_api_key()
            
            # Reinitialize the client with the loaded settings
            self.client = OpenAIClient(self.api_key)
            self.client.set_model(self.model)
            self.initialized = True
                
        except Exception as e:
            logger.error(f"Error loading OpenAI config: {str(e)}")
            # Even on error, create a client with the default model and loaded API key
            self.client = OpenAIClient(get_api_key())
            self.client.set_model(DEFAULT_MODEL)
            self.initialized = True
            
    def update_api_key(self, new_api_key):
        """Update the API key and save to config file"""
        if save_api_key(new_api_key):
            self.api_key = new_api_key
            self.client = OpenAIClient(self.api_key)
            self.client.set_model(self.model)
            return True
        return False
    
    def update_system_prompt(self, new_prompt):
        """Update the system prompt and save to config file"""
        self.system_prompt = new_prompt
        self.save_config()
        return True
            
    def save_config(self):
        """Save configuration to file"""
        config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config", "openai_config.json")
        try:
            # Create config directory if it doesn't exist
            config_dir = os.path.dirname(config_path)
            os.makedirs(config_dir, exist_ok=True)
            
            with open(config_path, 'w') as f:
                json.dump({
                    "model": self.model,
                    "system_prompt": self.system_prompt
                }, f, indent=2)
                
            # Update the client with new settings
            self.client = OpenAIClient(self.api_key)
            self.client.set_model(self.model)
        except Exception as e:
            logger.error(f"Error saving OpenAI config: {str(e)}")
            
    def generate_response(self, messages, conversation_files=None):
        """Generate response from OpenAI API"""
        try:
            if not self.api_key:
                return "Error: No API key provided. Please configure your OpenAI API key using the 'Add API Key' button."
            
            return self.client.chat_completion(messages, conversation_files)
        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            return f"Error calling OpenAI API: {str(e)}"

class ChatHistoryItem(wx.Panel):
    """A single item in the conversation history"""
    def __init__(self, parent, conversation, on_select, on_delete, on_rename):
        super().__init__(parent)
        self.conversation = conversation
        self.on_select = on_select
        self.on_delete = on_delete
        self.on_rename = on_rename
        
        self.SetBackgroundColour(wx.Colour(40, 40, 40))
        
        # Create a horizontal sizer
        sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # Add conversation title
        self.title_text = wx.StaticText(self, label=conversation.title)
        self.title_text.SetForegroundColour(wx.Colour(240, 240, 240))
        sizer.Add(self.title_text, 1, wx.EXPAND | wx.ALL, 5)
        
        # Create a button panel to ensure buttons stay together
        button_panel = wx.Panel(self)
        button_panel.SetBackgroundColour(wx.Colour(40, 40, 40))
        button_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # Add rename button
        rename_btn = wx.Button(button_panel, label="✏️", size=(30, 24))
        rename_btn.SetToolTip("Rename conversation")
        rename_btn.Bind(wx.EVT_BUTTON, self._on_rename)
        button_sizer.Add(rename_btn, 0, wx.ALL, 2)
        
        # Add delete button
        delete_btn = wx.Button(button_panel, label="×", size=(30, 24))
        delete_btn.SetToolTip("Delete this conversation")
        delete_btn.Bind(wx.EVT_BUTTON, self._on_delete)
        button_sizer.Add(delete_btn, 0, wx.ALL, 2)
        
        button_panel.SetSizer(button_sizer)
        sizer.Add(button_panel, 0, wx.ALIGN_CENTER_VERTICAL)
        
        # Make the whole panel clickable to select conversation
        self.Bind(wx.EVT_LEFT_DOWN, self._on_select_this)
        self.title_text.Bind(wx.EVT_LEFT_DOWN, self._on_select_this)
        
        self.SetSizer(sizer)
        
        # Set minimum size to ensure visibility
        self.SetMinSize(wx.Size(150, 30))
        
    def _on_select_this(self, event):
        """Handle click on panel to select this conversation"""
        try:
            conv_id = self.conversation.id if hasattr(self, 'conversation') else None
            
            if conv_id is None:
                logger.error("Cannot select conversation: No conversation ID found")
                return
                
            self.on_select(conv_id)
        except Exception as e:
            logger.error(f"Error selecting conversation {getattr(self.conversation, 'id', 'unknown')}: {str(e)}")
            # Use CallAfter to ensure UI operations happen in the main thread
            wx.CallAfter(lambda: wx.MessageBox(f"Error selecting conversation: {str(e)}", 
                                             "Error", wx.OK | wx.ICON_ERROR))
        
        # Don't skip the event to prevent bubbling up
        event.Skip(False)
        
    def update_title(self):
        """Update the displayed title to match the conversation title"""
        try:
            if not hasattr(self, 'conversation'):
                logger.error("Cannot update title: No conversation object found")
                return
                
            if not hasattr(self, 'title_text') or not self.title_text:
                logger.error("Cannot update title: No title_text control found")
                return
                
            current_title = getattr(self.conversation, 'title', 'Untitled')
            old_title = self.title_text.GetLabel()
            
            if current_title != old_title:
                logger.info(f"Updating conversation title from '{old_title}' to '{current_title}'")
                self.title_text.SetLabel(current_title)
                self.Layout()
            else:
                logger.debug(f"Title unchanged: '{current_title}'")
                
        except Exception as e:
            logger.error(f"Error updating title for conversation {getattr(self.conversation, 'id', 'unknown')}: {str(e)}")
        
    def _on_delete(self, event):
        """Handle delete"""
        try:
            # Store conversation ID in case we need it for error reporting
            conv_id = self.conversation.id if hasattr(self, 'conversation') else None
            
            if conv_id is None:
                logger.error("Cannot delete conversation: No conversation ID found")
                return
                
            # Call the callback with proper error handling
            self.on_delete(conv_id)
        except Exception as e:
            logger.error(f"Error triggering delete for conversation {getattr(self.conversation, 'id', 'unknown')}: {str(e)}")
            wx.CallAfter(lambda: wx.MessageBox(f"Error deleting conversation: {str(e)}", "Error", wx.OK | wx.ICON_ERROR))
        
        # Don't skip the event to prevent bubbling up
        
    def _on_rename(self, event):
        """Handle rename"""
        try:
            # Store conversation ID in case we need it for error reporting
            conv_id = self.conversation.id if hasattr(self, 'conversation') else None
            
            if conv_id is None:
                logger.error("Cannot rename conversation: No conversation ID found")
                return
                
            # Call the callback with proper error handling
            self.on_rename(conv_id)
        except Exception as e:
            logger.error(f"Error triggering rename for conversation {getattr(self.conversation, 'id', 'unknown')}: {str(e)}")
            wx.CallAfter(lambda: wx.MessageBox(f"Error renaming conversation: {str(e)}", "Error", wx.OK | wx.ICON_ERROR))
        
        # Don't skip the event to prevent bubbling up which may cause crashes

class SidebarPanel(wx.Panel):
    """Left sidebar panel for chat history"""
    def __init__(self, parent, conversation_manager, on_select_conversation):
        super().__init__(parent, style=wx.BORDER_NONE)
        
        self.conversation_manager = conversation_manager
        self.on_select_conversation = on_select_conversation
        self.openai_manager = parent.openai_manager  # Reference to the parent's OpenAI manager
        
        # Set background color
        self.SetBackgroundColour(wx.Colour(40, 40, 40))
        
        # Create main sizer
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Create new chat button
        self.new_chat_btn = wx.Button(self, label="+ New Chat")
        self.new_chat_btn.SetForegroundColour(wx.Colour(255, 255, 255))  # White text
        self.new_chat_btn.SetBackgroundColour(wx.Colour(0, 120, 215))  # Blue background for better visibility
        self.new_chat_btn.Bind(wx.EVT_BUTTON, self._on_new_chat)
        
        # Create a scrolled window for chat history with improved settings
        self.history_panel = wx.ScrolledWindow(self, style=wx.VSCROLL)
        self.history_panel.SetScrollRate(0, 10)
        self.history_panel.SetBackgroundColour(wx.Colour(40, 40, 40))
        
        # Set minimum size for history panel to ensure it's always visible
        self.history_panel.SetMinSize(wx.Size(200, 100))
        
        # Sizer for history items
        self.history_sizer = wx.BoxSizer(wx.VERTICAL)
        self.history_panel.SetSizer(self.history_sizer)
        
        # Add elements to the main sizer
        main_sizer.Add(self.new_chat_btn, 0, wx.EXPAND | wx.ALL, 5)
        main_sizer.Add(wx.StaticLine(self), 0, wx.EXPAND | wx.ALL, 2)
        main_sizer.Add(self.history_panel, 1, wx.EXPAND)
        
        self.SetSizer(main_sizer)
        
        # Populate history
        self.populate_history()
        
    def _on_new_chat(self, event):
        """Handle new chat button click"""
        try:
            # Create a new conversation with the OpenAIManager's system prompt
            system_message = self.openai_manager.system_prompt
            logger.info(f"Creating new chat with system message: {system_message[:50]}...")
            
            # Create the new conversation
            conversation = self.conversation_manager.new_conversation(system_message)
            logger.info(f"Created new conversation with ID: {conversation.id}")
            
            # Update the UI and ensure proper layout
            self.populate_history()
            
            # Notify the main panel
            self.on_select_conversation(conversation.id)
            
            # Extra step to ensure all conversations are visible
            wx.CallAfter(self._ensure_history_visible)
            
            # Log status for debugging
            logger.info(f"Current conversations: {len(self.conversation_manager.conversations)}")
            logger.info(f"Current conversation ID: {self.conversation_manager.current_conversation_id}")
            
        except Exception as e:
            logger.error(f"Error creating new chat: {str(e)}")
            wx.MessageBox(f"Error creating new chat: {str(e)}", "Error", wx.OK | wx.ICON_ERROR)
            
    def _ensure_history_visible(self):
        """Additional method to ensure history items are visible after updates"""
        try:
            # Force layout update
            self.history_panel.FitInside()
            self.history_panel.Layout()
            self.Layout()
            
            # Force redraw
            self.history_panel.Refresh()
            self.Refresh()
            
            # Make sure first conversation is visible
            if self.history_sizer.GetItemCount() > 0:
                self.history_panel.ScrollChildIntoView(self.history_sizer.GetItem(0).GetWindow())
        except Exception as e:
            logger.error(f"Error ensuring history visibility: {str(e)}")
        
    def _on_select_conversation(self, conversation_id):
        """Handle conversation selection"""
        self.on_select_conversation(conversation_id)
        
    def _on_delete_conversation(self, conversation_id):
        """Handle conversation deletion"""
        try:
            # Verify the conversation exists
            if conversation_id not in self.conversation_manager.conversations:
                logger.warning(f"Attempted to delete non-existent conversation: {conversation_id}")
                return
                
            # Show confirmation dialog
            dlg = wx.MessageDialog(
                self, 
                "Are you sure you want to delete this conversation?", 
                "Confirm Deletion",
                wx.YES_NO | wx.ICON_QUESTION
            )
            
            try:
                if dlg.ShowModal() == wx.ID_YES:
                    # Store the current ID before deletion
                    current_id_before_delete = self.conversation_manager.current_conversation_id
                    
                    # Delete the conversation
                    result = self.conversation_manager.delete_conversation(conversation_id)
                    if not result:
                        logger.warning(f"Failed to delete conversation {conversation_id}")
                        wx.MessageBox("Failed to delete conversation.", "Warning", wx.OK | wx.ICON_WARNING)
                        return
                    
                    # Update the UI
                    self.populate_history()
                    
                    # Select the current conversation if it changed
                    if self.conversation_manager.current_conversation_id != current_id_before_delete:
                        logger.info(f"Switching to conversation {self.conversation_manager.current_conversation_id} after deletion")
                        self.on_select_conversation(self.conversation_manager.current_conversation_id)
            finally:
                # Make sure dialog is always destroyed
                dlg.Destroy()
                
        except Exception as e:
            logger.error(f"Error in _on_delete_conversation: {str(e)}")
            wx.MessageBox(f"Error deleting conversation: {str(e)}", "Error", wx.OK | wx.ICON_ERROR)
            
            # Make sure we have a valid conversation after error
            if not self.conversation_manager.conversations:
                try:
                    new_conversation = self.conversation_manager.new_conversation()
                    self.populate_history()
                    self.on_select_conversation(new_conversation.id)
                    logger.info(f"Created new conversation {new_conversation.id} after deletion error recovery")
                except Exception as recovery_error:
                    logger.error(f"Failed to recover from deletion error: {str(recovery_error)}")
                    wx.MessageBox("Critical error in conversation management. Please restart the application.", 
                                 "Critical Error", wx.OK | wx.ICON_ERROR)
    
    def _on_rename_conversation(self, conversation_id):
        """Handle conversation rename"""
        try:
            # Verify the conversation exists
            if conversation_id not in self.conversation_manager.conversations:
                logger.warning(f"Attempted to rename non-existent conversation: {conversation_id}")
                return
                
            # Get the current conversation
            conversation = self.conversation_manager.conversations.get(conversation_id)
            
            # Create rename dialog
            dlg = wx.TextEntryDialog(
                self,
                "Enter a new name for this conversation:",
                "Rename Conversation",
                conversation.title
            )
            
            try:
                if dlg.ShowModal() == wx.ID_OK:
                    new_title = dlg.GetValue().strip()
                    if new_title:
                        # Rename the conversation
                        result = self.conversation_manager.rename_conversation(conversation_id, new_title)
                        if not result:
                            logger.warning(f"Failed to rename conversation {conversation_id}")
                            wx.MessageBox("Failed to rename conversation.", "Warning", wx.OK | wx.ICON_WARNING)
                            return
                            
                        # Update the UI - ensure this happens in the main thread
                        wx.CallAfter(self.populate_history)
            finally:
                # Make sure dialog is always destroyed
                dlg.Destroy()
                
        except Exception as e:
            logger.error(f"Error in _on_rename_conversation: {str(e)}")
            wx.MessageBox(f"Error renaming conversation: {str(e)}", "Error", wx.OK | wx.ICON_ERROR)
        
    def populate_history(self):
        """Populate the history panel with conversation items"""
        try:
            # Clear existing items
            self.history_sizer.Clear(True)
            
            # Get all conversations sorted by last updated time (newest first)
            sorted_conversations = sorted(
                self.conversation_manager.conversations.values(),
                key=lambda x: x.last_updated,
                reverse=True
            )
            
            # Add history items
            for conversation in sorted_conversations:
                item = ChatHistoryItem(
                    self.history_panel, 
                    conversation,
                    self._on_select_conversation,
                    self._on_delete_conversation,
                    self._on_rename_conversation
                )
                
                # Highlight the current conversation
                if conversation.id == self.conversation_manager.current_conversation_id:
                    item.SetBackgroundColour(wx.Colour(60, 60, 60))
                    
                self.history_sizer.Add(item, 0, wx.EXPAND | wx.BOTTOM, 2)
                
            # Refresh the layout - make sure to update the scrollbars
            self.history_panel.FitInside()
            self.history_panel.Layout()
            self.Layout()
            
            # Force a redraw to ensure all items are visible
            self.history_panel.Refresh()
            self.Refresh()
        except Exception as e:
            logger.error(f"Error in populate_history: {str(e)}")
            wx.MessageBox(f"Error updating chat history: {str(e)}", "Error", wx.OK | wx.ICON_ERROR)
        
    def update_current_conversation_title(self):
        """Update the title of the current conversation in the UI"""
        try:
            if not self.conversation_manager.current_conversation_id:
                logger.warning("No current conversation ID, can't update title")
                return
                
            current_id = self.conversation_manager.current_conversation_id
            conversation = self.conversation_manager.conversations.get(current_id)
            
            if not conversation:
                logger.warning(f"Conversation not found for ID: {current_id}")
                return
                
            logger.info(f"Updating conversation title in UI: ID={current_id}, title='{conversation.title}'")
                
            # Use wx.CallAfter to ensure UI updates happen in the main thread
            wx.CallAfter(self._update_conversation_title_ui)
            # Also refresh the entire history to ensure all titles are up to date
            wx.CallAfter(self.populate_history)
        except Exception as e:
            logger.error(f"Error in update_current_conversation_title: {str(e)}")
            # Don't show message box here as it's a background operation
            
    def _update_conversation_title_ui(self):
        """Internal method to update the UI with new title - must be called from main thread"""
        try:
            # Find all child windows in the history panel
            children = self.history_panel.GetChildren()
            current_id = self.conversation_manager.current_conversation_id
            
            if not current_id:
                logger.warning("No current conversation ID in _update_conversation_title_ui")
                return
                
            # Get the current conversation title
            current_conv = self.conversation_manager.conversations.get(current_id)
            if not current_conv:
                logger.warning(f"Conversation {current_id} not found in _update_conversation_title_ui")
                return
                
            logger.info(f"Updating UI title for conversation {current_id}: '{current_conv.title}'")
            
            # Update the title of the matching conversation item
            found = False
            for child in children:
                if isinstance(child, ChatHistoryItem) and child.conversation.id == current_id:
                    logger.info(f"Found matching ChatHistoryItem, updating title to: '{current_conv.title}'")
                    child.update_title()
                    found = True
                    break
                    
            if not found:
                logger.warning(f"No matching ChatHistoryItem found for conversation {current_id}")
                # Force a refresh of the whole history panel
                self.populate_history()
        except Exception as e:
            logger.error(f"Error updating conversation title UI: {str(e)}")
            # Try to repopulate the history as a fallback
            try:
                self.populate_history()
            except Exception as repop_error:
                logger.error(f"Also failed to repopulate history: {str(repop_error)}")

class FileDragAndDropHelper:
    """Helper class for drag and drop file operations"""
    
    @staticmethod
    def highlight_drop_target(panel, highlight=True):
        """Highlight or unhighlight a panel as a drop target"""
        if highlight:
            panel.SetBackgroundColour(wx.Colour(120, 180, 240))  # Medium blue for better visibility
        else:
            panel.SetBackgroundColour(wx.Colour(180, 180, 180))  # Restore to darker gray
        panel.Refresh()

class EnhancedFileDropTarget(wx.FileDropTarget):
    """Enhanced drag and drop handling for files with visual feedback"""
    def __init__(self, chat_panel, drop_area=None):
        super(EnhancedFileDropTarget, self).__init__()
        self.chat_panel = chat_panel
        self.drop_area = drop_area  # The specific panel to highlight
        
    def OnDragOver(self, x, y, defResult):
        """Show visual feedback during drag"""
        # Call base class method to get the default result
        result = super(EnhancedFileDropTarget, self).OnDragOver(x, y, defResult)
        
        # Highlight the drop area if specified
        if self.drop_area:
            FileDragAndDropHelper.highlight_drop_target(self.drop_area, True)
            
        return result
    
    def OnLeave(self):
        """Reset visual feedback when drag leaves the drop target"""
        if self.drop_area:
            FileDragAndDropHelper.highlight_drop_target(self.drop_area, False)
        return super(EnhancedFileDropTarget, self).OnLeave()
        
    def OnDropFiles(self, x, y, filenames):
        """Handle files dropped onto the window"""
        # Reset visual feedback
        if self.drop_area:
            FileDragAndDropHelper.highlight_drop_target(self.drop_area, False)
            
        if filenames and len(filenames) > 0:
            # Only process the first file
            try:
                # Clear any existing file first
                if self.chat_panel.current_file:
                    self.chat_panel.on_clear_file(None)
                
                # Process the dropped file
                self.chat_panel.process_file(filenames[0])
                
                # Set focus to input field for convenient typing after drop
                self.chat_panel.input_field.SetFocus()
                return True
            except Exception as e:
                logger.error(f"Error handling dropped file: {str(e)}")
                wx.MessageBox(f"Error processing dropped file: {str(e)}", "Error", wx.OK | wx.ICON_ERROR)
        return False

class ChatPanel(wx.Panel):
    """Chat UI panel implementation"""
    
    def __init__(self, parent):
        wx.Panel.__init__(self, parent, id=wx.ID_ANY)
        
        self.openai_manager = OpenAIManager()
        # Ensure OpenAI Manager is fully initialized
        self.openai_manager.load_config()
        self.conversation_manager = ConversationManager()
        
        # Check for optional dependencies
        self.has_pypdf2 = self._check_pypdf2_installed()
        self.has_python_docx = self._check_python_docx_installed()
        
        # Create system message for new conversations
        self.default_system_message = self.openai_manager.system_prompt
        
        # Ensure there's at least one conversation
        if not self.conversation_manager.conversations:
            self.conversation_manager.new_conversation(self.default_system_message)
            
        # Create a horizontal sizer for the sidebar and chat area
        main_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # Create the sidebar
        self.sidebar = SidebarPanel(
            self, 
            self.conversation_manager,
            self._on_select_conversation
        )
        
        # Create the chat area container
        chat_container = wx.Panel(self)
        chat_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Create the chat components
        self.chat_history = wx.TextCtrl(chat_container, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_RICH2)
        self.chat_history.SetBackgroundColour(wx.Colour(240, 240, 240))
        self.chat_history.SetMinSize(wx.Size(300, 200))
        
        # Input area with file upload button
        input_panel = wx.Panel(chat_container)
        input_panel.SetBackgroundColour(wx.Colour(245, 245, 245))
        input_panel_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Message input field
        self.input_field = wx.TextCtrl(input_panel, style=wx.TE_MULTILINE)
        self.input_field.SetMinSize(wx.Size(-1, 60))  # Taller input field
        self.input_field.SetHint("Type your message here...")
        self.input_field.Bind(wx.EVT_KEY_DOWN, self.on_input_key_down)
        
        # Controls area with send button and file upload
        controls_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # File drop area panel that spans most of the width
        file_drop_container = wx.Panel(input_panel)
        file_drop_container.SetBackgroundColour(wx.Colour(180, 180, 180))  # Darker gray for better contrast
        file_drop_container_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # Left side with file upload button and drop label
        self.file_drop_panel = wx.Panel(file_drop_container)
        self.file_drop_panel.SetBackgroundColour(wx.Colour(180, 180, 180))  # Match parent color
        file_panel_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # File upload button with explicit text color and border
        file_button = wx.Button(self.file_drop_panel, label="📎 Attach")
        file_button.SetForegroundColour(wx.Colour(0, 0, 0))  # Black text
        file_button.SetToolTip("Attach a file")
        file_button.Bind(wx.EVT_BUTTON, self.on_file_upload)
        file_panel_sizer.Add(file_button, 0, wx.ALIGN_CENTER_VERTICAL | wx.ALL, 2)
        
        # Drop zone indicator with better contrast
        drop_label = wx.StaticText(self.file_drop_panel, label="or drop files here")
        drop_label.SetForegroundColour(wx.Colour(50, 50, 50))  # Darker text for better contrast
        file_panel_sizer.Add(drop_label, 0, wx.ALIGN_CENTER_VERTICAL | wx.ALL, 2)
        
        self.file_drop_panel.SetSizer(file_panel_sizer)
        
        # Right side with file indicator (part of the gray panel)
        self.file_indicator_panel = wx.Panel(file_drop_container)
        self.file_indicator_panel.SetBackgroundColour(wx.Colour(180, 180, 180))  # Match parent color
        file_indicator_sizer = wx.BoxSizer(wx.HORIZONTAL)
        
        # Attached file indicator with better contrast
        self.file_indicator = wx.StaticText(self.file_indicator_panel, label="")
        self.file_indicator.SetForegroundColour(wx.Colour(0, 100, 0))  # Dark green for file name
        file_indicator_sizer.Add(self.file_indicator, 1, wx.ALIGN_CENTER_VERTICAL | wx.LEFT, 5)
        
        # Clear file button (hidden initially) - with black text
        self.clear_file_btn = wx.Button(self.file_indicator_panel, label="✖", size=(24, 24))
        self.clear_file_btn.SetForegroundColour(wx.Colour(0, 0, 0))  # Black text for contrast
        self.clear_file_btn.SetToolTip("Clear attached file")
        self.clear_file_btn.Bind(wx.EVT_BUTTON, self.on_clear_file)
        self.clear_file_btn.Hide()
        file_indicator_sizer.Add(self.clear_file_btn, 0, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 5)
        
        self.file_indicator_panel.SetSizer(file_indicator_sizer)
        
        # Add both panels to the container
        file_drop_container_sizer.Add(self.file_drop_panel, 0, wx.ALIGN_CENTER_VERTICAL)
        file_drop_container_sizer.Add(self.file_indicator_panel, 1, wx.EXPAND | wx.ALIGN_CENTER_VERTICAL)
        file_drop_container.SetSizer(file_drop_container_sizer)
        
        # Make panels drop targets with visual feedback
        self.file_drop_panel.SetDropTarget(EnhancedFileDropTarget(self, self.file_drop_panel))
        self.file_indicator_panel.SetDropTarget(EnhancedFileDropTarget(self, self.file_drop_panel))
        file_drop_container.SetDropTarget(EnhancedFileDropTarget(self, self.file_drop_panel))
        input_panel.SetDropTarget(EnhancedFileDropTarget(self, self.file_drop_panel))
        self.chat_history.SetDropTarget(EnhancedFileDropTarget(self, self.file_drop_panel))
        
        # Add file drop container to controls sizer - NO right margin so it touches the send button
        controls_sizer.Add(file_drop_container, 1, wx.EXPAND)
        
        # Send button panel with matching gray background
        send_button_panel = wx.Panel(input_panel)
        send_button_panel.SetBackgroundColour(wx.Colour(180, 180, 180))  # Match file drop area color
        send_button_panel_sizer = wx.BoxSizer(wx.VERTICAL)
        
        # Send button with explicit text color
        send_button = wx.Button(send_button_panel, label="Send")
        send_button.SetForegroundColour(wx.Colour(0, 0, 0))  # Black text
        
        # Make sure we're binding the send button correctly
        logger.info("Binding Send button to on_send method")
        send_button.Bind(wx.EVT_BUTTON, self.on_send)
        
        # Center the button vertically in its panel
        send_button_panel_sizer.AddStretchSpacer()
        send_button_panel_sizer.Add(send_button, 0, wx.ALIGN_CENTER | wx.ALL, 2)
        send_button_panel_sizer.AddStretchSpacer()
        
        send_button_panel.SetSizer(send_button_panel_sizer)
        controls_sizer.Add(send_button_panel, 0, wx.EXPAND)
        
        # Add components to input panel
        input_panel_sizer.Add(self.input_field, 1, wx.EXPAND | wx.BOTTOM, 5)
        input_panel_sizer.Add(controls_sizer, 0, wx.EXPAND)
        input_panel.SetSizer(input_panel_sizer)
        
        # Model selector with white text label (for dark background)
        model_sizer = wx.BoxSizer(wx.HORIZONTAL)
        model_label = wx.StaticText(chat_container, label="Model:")
        model_label.SetForegroundColour(wx.Colour(255, 255, 255))  # White text for dark theme
        self.model_selector = wx.Choice(chat_container, choices=["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"])
        self.model_selector.SetSelection(0)  # Default to gpt-3.5-turbo
        self.model_selector.Bind(wx.EVT_CHOICE, self.on_model_change)
        model_sizer.Add(model_label, 0, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 5)
        model_sizer.Add(self.model_selector, 0, wx.ALIGN_CENTER_VERTICAL | wx.RIGHT, 15)
        
        # Add API Key button
        api_key_btn = wx.Button(chat_container, label="Add API Key")
        api_key_btn.SetToolTip("Configure your OpenAI API key")
        api_key_btn.Bind(wx.EVT_BUTTON, self.on_configure_api_key)
        model_sizer.Add(api_key_btn, 0, wx.ALIGN_CENTER_VERTICAL)
        
        # Set background color for model selector area
        chat_container.SetBackgroundColour(wx.Colour(60, 60, 60))  # Dark background for model selector
        
        # Combine all in the chat sizer
        chat_sizer.Add(self.chat_history, 1, wx.EXPAND | wx.ALL, 5)
        chat_sizer.Add(input_panel, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 5)
        chat_sizer.Add(model_sizer, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 5)
        
        chat_container.SetSizer(chat_sizer)
        
        # Add the sidebar and chat area to the main sizer
        main_sizer.Add(self.sidebar, 0, wx.EXPAND)
        main_sizer.Add(chat_container, 1, wx.EXPAND)
        
        self.SetSizer(main_sizer)
        
        # Initialize file upload variables
        self.current_file = None
        
        # Load the current conversation
        self._load_current_conversation()
        
    def _check_pypdf2_installed(self):
        """Check if PyPDF2 is installed"""
        try:
            import PyPDF2
            logger.info("PyPDF2 is installed")
            return True
        except ImportError:
            logger.warning("PyPDF2 is not installed - PDF text extraction will not be available")
            return False
            
    def _check_python_docx_installed(self):
        """Check if python-docx is installed"""
        try:
            import docx
            logger.info("python-docx is installed")
            return True
        except ImportError:
            logger.warning("python-docx is not installed - DOCX text extraction will not be available")
            return False
        
    def _on_select_conversation(self, conversation_id):
        """Handle conversation selection from sidebar"""
        # Switch to the selected conversation
        conversation = self.conversation_manager.switch_conversation(conversation_id)
        
        # Update the UI
        self.sidebar.populate_history()
        
        # Load the conversation history
        self._load_current_conversation()
        
    def _load_current_conversation(self):
        """Load the current conversation into the chat history"""
        # Clear the chat history
        self.chat_history.Clear()
        
        # Get the current conversation
        conversation = self.conversation_manager.get_current_conversation()
        
        # Display messages
        for msg in conversation.messages:
            if msg["role"] == "system":
                continue  # Skip system messages
            elif msg["role"] == "user":
                self.add_user_message(msg["content"], add_to_conversation=False)
            elif msg["role"] == "assistant":
                self.add_assistant_message(msg["content"], add_to_conversation=False)
                
        # If no messages yet, show the welcome message
        if not [msg for msg in conversation.messages if msg["role"] != "system"]:
            welcome_message = "Welcome to KiCad AI Plugin by ALT TAB! How can I help you today?"
            
            # If API key is not set, add instructions
            if not self.openai_manager.api_key:
                welcome_message += "\n\nPlease configure your OpenAI API key by clicking the 'Add API Key' button before sending messages."
                
            self.add_assistant_message(welcome_message, add_to_conversation=False)
        
    def on_model_change(self, event):
        """Handle model selection change"""
        model = self.model_selector.GetStringSelection()
        self.openai_manager.model = model
        self.openai_manager.client.set_model(model)
        self.openai_manager.save_config()
        self.add_status_message(f"Model changed to {model}")
        
    def on_send(self, event):
        """Handle sending a message"""
        # Add debug logging
        logger.info("on_send method called")
        
        message = self.input_field.GetValue().strip()
        logger.info(f"Message to send: '{message}'")
        
        # Check if we have a message or a file
        if not message and not self.current_file:
            logger.info("No message or file to send")
            return
            
        # Check if API key is configured
        if not self.openai_manager.api_key:
            self.add_assistant_message(
                "Please configure your OpenAI API key first by clicking the 'Add API Key' button.",
                add_to_conversation=False
            )
            return
            
        # Initialize command_processor attribute if it doesn't exist yet
        if not hasattr(self, 'command_processor'):
            self.command_processor = None

        # Track if this is the first message in conversation
        conversation = self.conversation_manager.get_current_conversation()
        is_first_message = len([m for m in conversation.messages if m["role"] == "user"]) == 0
            
        # Process message with file if present
        if self.current_file:
            logger.info(f"Sending message with attached file: {self.current_file['name']}")
            self.add_user_message_with_file(message, self.current_file)
            # Clear the file after sending
            self.on_clear_file(None)
        else:
            logger.info("Sending text message")
            self.add_user_message(message)
            
        self.input_field.SetValue("")
        
        # Show typing indicator
        self.add_status_message("Thinking...")
        
        # Update sidebar if this was the first message (to show the new title)
        if is_first_message:
            wx.CallAfter(self.sidebar.update_current_conversation_title)
            wx.CallAfter(self.sidebar.populate_history)
        
        # Process in a separate thread to keep UI responsive
        threading.Thread(target=self.process_with_ai, daemon=True).start()
        
    def process_with_ai(self):
        """Process the message with OpenAI in a background thread"""
        try:
            # Get current conversation
            conversation = self.conversation_manager.get_current_conversation()
            
            # Log the conversation files for debugging
            logger.info(f"Processing with conversation files: {len(conversation.files) if hasattr(conversation, 'files') else 0} files")
            
            # Regular chat processing
            response = self.openai_manager.generate_response(
                conversation.messages,
                conversation.files
            )
            
            # Add to conversation history
            conversation.add_message("assistant", response)
            
            # Save conversations
            self.conversation_manager.save_conversations()
            
            # Update UI
            wx.CallAfter(self.sidebar.update_current_conversation_title)
            wx.CallAfter(self.sidebar.populate_history)
            
            # Remove typing indicator and add response
            wx.CallAfter(self.remove_status_message)
            wx.CallAfter(self.add_assistant_message, response, add_to_conversation=False)
            
        except Exception as e:
            logger.error(f"Error in AI processing: {str(e)}")
            wx.CallAfter(self.remove_status_message)
            wx.CallAfter(self.add_assistant_message, f"Error processing your request: {str(e)}", add_to_conversation=False)
            
    def on_settings(self, event):
        """Show settings dialog"""
        # We no longer need a full settings dialog since we're using hardcoded API key
        dlg = wx.MessageDialog(self, "This plugin uses a pre-configured API key.\nYou can change the model using the dropdown menu.", "AI Settings", wx.OK | wx.ICON_INFORMATION)
        dlg.ShowModal()
        dlg.Destroy()
            
    def add_user_message(self, message, add_to_conversation=True):
        """Add a user message to the chat history"""
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.BLACK))
        self.chat_history.AppendText("You: ")
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.BLACK))
        self.chat_history.AppendText(f"{message}\n")
        
        # Add to conversation if needed
        if add_to_conversation:
            conversation = self.conversation_manager.get_current_conversation()
            previous_user_msgs = len([m for m in conversation.messages if m["role"] == "user"])
            conversation.add_message("user", message)
            self.conversation_manager.save_conversations()
            
            # If this is the first user message, update the UI with the new title
            if previous_user_msgs == 0:
                logger.info("First user message added, updating title in sidebar")
                self.sidebar.update_current_conversation_title()
                wx.CallAfter(self.sidebar.populate_history)
        
    def add_assistant_message(self, message, add_to_conversation=True):
        """Add an assistant message to the chat history"""
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.Colour(0, 100, 0)))
        self.chat_history.AppendText("Assistant: ")
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.BLACK))
        self.chat_history.AppendText(f"{message}\n")
        
        # Add to conversation if needed
        if add_to_conversation:
            conversation = self.conversation_manager.get_current_conversation()
            conversation.add_message("assistant", message)
            self.conversation_manager.save_conversations()
        
    def add_status_message(self, message):
        """Add a status message to the chat history"""
        self.status_message_pos = self.chat_history.GetLastPosition()
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.BLUE))
        self.chat_history.AppendText(f"[{message}]\n")
        
    def remove_status_message(self):
        """Remove the last status message"""
        if hasattr(self, 'status_message_pos'):
            last_pos = self.chat_history.GetLastPosition()
            self.chat_history.Remove(self.status_message_pos, last_pos)

    def on_file_upload(self, event):
        """Handle file upload button click"""
        wildcard = "All files (*.*)|*.*"
        dlg = wx.FileDialog(
            self, message="Choose a file to upload",
            defaultDir=os.getcwd(),
            defaultFile="",
            wildcard=wildcard,
            style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST
        )
        
        if dlg.ShowModal() == wx.ID_OK:
            path = dlg.GetPath()
            self.process_file(path)
            
        dlg.Destroy()
        
    def on_clear_file(self, event):
        """Clear the currently attached file"""
        self.current_file = None
        self.file_indicator.SetLabel("")
        self.clear_file_btn.Hide()
        self.file_indicator_panel.Layout()
        self.Layout()
        
    def process_file(self, filepath):
        """Process a file upload"""
        try:
            # Get file info
            filename = os.path.basename(filepath)
            filesize = os.path.getsize(filepath)
            file_ext = os.path.splitext(filename)[1].lower()
            
            # Check file size
            if filesize > MAX_FILE_SIZE:
                self.add_status_message(f"File is too large ({filesize/1024/1024:.1f} MB). Maximum size is 5 MB.")
                return False
                
            # Determine file type and process accordingly
            mime_type = self._get_mime_type(file_ext)
            is_text_file = mime_type.startswith('text/') or file_ext in ['.md', '.py', '.js', '.html', '.css', '.json', '.yml', '.yaml']
            
            # Check if this is a KiCad file
            try:
                # Use the imported is_kicad_file function if available
                kicad_file = is_kicad_file(file_ext)
            except:
                # Fallback in case of import error
                kicad_file = file_ext.lower() in ['.kicad_pcb', '.kicad_sch', '.kicad_pro', '.net', 
                                                 '.lib', '.kicad_mod', '.kicad_wks', '.kicad_sym']
            
            file_data = {
                        "name": filename,
                "path": filepath,
                "type": mime_type,
                "size": filesize,
                "is_binary": not (is_text_file or kicad_file),  # KiCad files are treated as text
                "mime_type": mime_type
            }
            
            # Add KiCad specific info if applicable
            if kicad_file:
                file_data["is_kicad_file"] = True
                file_data["kicad_file_type"] = file_ext[1:]  # Store the KiCad file type without the dot
            
            # Process file content based on type
            try:
                # Process text files (including KiCad files)
                if is_text_file or kicad_file:
                    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                        text_content = f.read()
                        file_data["content"] = text_content
                        
                        # For API compatibility, also add a content preview
                        content_preview = text_content[:4000]
                        if len(text_content) > 4000:
                            content_preview += "...(truncated)"
                        file_data["content_preview"] = content_preview
                        
                        # Extract KiCad specific information if applicable
                        if kicad_file:
                            try:
                                # Use the imported extract_kicad_file_info function if available
                                file_data["kicad_info"] = extract_kicad_file_info(filepath, file_ext, text_content)
                            except:
                                # Fallback in case of import error
                                file_data["kicad_info"] = {
                                    "file_type": f"KiCad {file_ext[1:].upper()} File",
                                    "summary": f"KiCad {file_ext[1:].upper()} File (detailed info not available)"
                                }
                
                # Process PDF files
                elif file_ext.lower() == '.pdf':
                    # Try various methods to extract PDF text
                    extracted_text = self._extract_pdf_text(filepath)
                    
                    # If we got no text extraction, attempt to use alternative methods
                    if "[PDF text extraction not available" in extracted_text:
                        # Try alternative pdf text extraction
                        alternative_text = self._extract_pdf_text_alternative(filepath)
                        if alternative_text:
                            extracted_text = alternative_text
                    
                    # If we still have no text, provide basic info
                    if not extracted_text or not extracted_text.strip() or "[PDF text extraction not available" in extracted_text:
                        # Fall back to providing some basic pdf info
                        basic_info = self._get_basic_pdf_info(filepath)
                        file_data["content"] = basic_info
                        file_data["content_preview"] = basic_info
                    else:
                        # We got text, use it
                        file_data["content"] = extracted_text
                        preview = extracted_text[:4000]
                        if len(extracted_text) > 4000:
                            preview += "...(truncated)"
                        file_data["content_preview"] = preview
                        
                    # Always provide the base64 encoded PDF for potential future use
                    with open(filepath, 'rb') as f:
                        binary_content = f.read()
                        file_data["content_b64"] = base64.b64encode(binary_content).decode('utf-8')
                
                # Process DOCX files
                elif file_ext.lower() == '.docx':
                    extracted_text = self._extract_docx_text(filepath)
                    if extracted_text and extracted_text.strip() and not extracted_text.startswith("[DOCX text extraction not available"):
                        file_data["content"] = extracted_text
                        preview = extracted_text[:4000]
                        if len(extracted_text) > 4000:
                            preview += "...(truncated)"
                        file_data["content_preview"] = preview
                    else:
                        # If text extraction fails, handle as binary
                        with open(filepath, 'rb') as f:
                            binary_content = f.read()
                            file_data["content_b64"] = base64.b64encode(binary_content).decode('utf-8')
                            file_data["content"] = f"[Binary DOCX file: {filename}, size: {filesize/1024:.1f} KB]"
                            file_data["content_preview"] = f"[Binary DOCX file: {filename}, size: {filesize/1024:.1f} KB]"
                
                # Process all other binary files
                else:
                    with open(filepath, 'rb') as f:
                        binary_content = f.read()
                        # Always encode binary files as base64 for API compatibility
                        file_data["content_b64"] = base64.b64encode(binary_content).decode('utf-8')
                        # Add a preview description
                        file_data["content"] = f"[Binary file: {filename}, size: {filesize/1024:.1f} KB]"
                        file_data["content_preview"] = f"[Binary file: {filename}, size: {filesize/1024:.1f} KB]"
                
            except UnicodeDecodeError:
                # Fall back to binary processing if text decoding fails
                with open(filepath, 'rb') as f:
                    binary_content = f.read()
                    file_data["content_b64"] = base64.b64encode(binary_content).decode('utf-8')
                    file_data["content"] = f"[Binary file: {filename}, size: {filesize/1024:.1f} KB]"
                    file_data["content_preview"] = f"[Binary file: {filename}, size: {filesize/1024:.1f} KB]"
                    file_data["is_binary"] = True
            
            # Store the current file for attachment
            self.current_file = file_data
            
            # Set the file indicator in the UI to show full filename
            self.file_indicator.SetLabel(f"Attached file: {filename}")
            self.clear_file_btn.Show()
            self.file_indicator_panel.Layout()
            
            # Show success indicator
            self._show_file_drop_success()
            
            return True
            
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            self.add_status_message(f"Error processing file: {str(e)}")
            return False
            
    def _extract_kicad_file_info(self, filepath, file_ext, content):
        """Extract relevant information from KiCad files"""
        try:
            info = {}
            
            # Process based on file type
            if file_ext == '.kicad_pcb':
                # Extract PCB information
                info["file_type"] = "KiCad PCB Layout"
                
                # Count number of layers
                import re
                layers_match = re.search(r'\(layers\s+([^)]+)\)', content)
                if layers_match:
                    layers_text = layers_match.group(1)
                    layers_count = len(re.findall(r'\([0-9]+\s+"[^"]+"\s+[^)]+\)', layers_text))
                    info["layers_count"] = layers_count
                
                # Count footprints
                footprints_count = content.count('(footprint ')
                info["footprints_count"] = footprints_count
                
                # Extract board dimensions if available
                edge_cuts_pattern = r'\(gr_rect\s+\(start\s+([\d\.-]+)\s+([\d\.-]+)\)\s+\(end\s+([\d\.-]+)\s+([\d\.-]+)\)'
                edge_matches = re.findall(edge_cuts_pattern, content)
                if edge_matches:
                    for match in edge_matches:
                        x1, y1, x2, y2 = map(float, match)
                        width = abs(x2 - x1)
                        height = abs(y2 - y1)
                        if "dimensions" not in info:
                            info["dimensions"] = []
                        info["dimensions"].append(f"{width:.2f}mm x {height:.2f}mm")
                
            elif file_ext == '.kicad_sch':
                # Extract schematic information
                info["file_type"] = "KiCad Schematic"
                
                # Count symbols (components)
                symbols_count = content.count('(symbol ')
                info["symbols_count"] = symbols_count
                
                # Check if JSON-based format (KiCad 6+)
                if content.lstrip().startswith('{'):
                    try:
                        import json
                        sch_data = json.loads(content)
                        if "sheets" in sch_data:
                            info["sheets_count"] = len(sch_data["sheets"])
                    except:
                        pass
                else:
                    # S-expression format - count sheets
                    sheets_count = content.count('(sheet ')
                    info["sheets_count"] = sheets_count
                
            elif file_ext == '.kicad_pro':
                # Extract project information
                info["file_type"] = "KiCad Project"
                
                # Try to parse as JSON (KiCad 6+)
                try:
                    import json
                    project_data = json.loads(content)
                    
                    # Extract version
                    if "version" in project_data:
                        info["kicad_version"] = project_data["version"]
                    
                    # Extract nets count if available
                    if "board" in project_data and "design_settings" in project_data["board"]:
                        settings = project_data["board"]["design_settings"]
                        if "rules" in settings and "netclass_patterns" in settings["rules"]:
                            info["netclasses"] = len(settings["rules"]["netclass_patterns"])
                except:
                    # Not JSON or parsing failed
                    pass
                
            elif file_ext == '.net':
                # Extract netlist information
                info["file_type"] = "KiCad Netlist"
                
                # Count components and nets
                components_count = content.count('(comp ')
                info["components_count"] = components_count
                
                nets_count = content.count('(net ')
                info["nets_count"] = nets_count
                
            elif file_ext == '.lib' or file_ext == '.kicad_sym':
                # Extract library information
                info["file_type"] = "KiCad Symbol Library"
                
                # Count symbols in the library
                if file_ext == '.kicad_sym':
                    symbols_count = content.count('(symbol ')
                    info["symbols_count"] = symbols_count
                else:
                    # Older .lib format
                    symbols_count = content.count('DEF ')
                    info["symbols_count"] = symbols_count
                
            elif file_ext == '.kicad_mod':
                # Extract module information
                info["file_type"] = "KiCad Footprint"
                
                # Try to get the footprint name
                import re
                name_match = re.search(r'\(footprint\s+"([^"]+)"', content)
                if name_match:
                    info["footprint_name"] = name_match.group(1)
                
                # Count pads
                pads_count = content.count('(pad ')
                info["pads_count"] = pads_count
                
            # Add a summary string that can be included in the prompt
            summary_parts = []
            for key, value in info.items():
                if key != "file_type":  # Skip file_type as it's added first
                    summary_parts.append(f"{key}: {value}")
                    
            info["summary"] = f"{info.get('file_type', 'KiCad File')}: " + ", ".join(summary_parts)
            
            return info
            
        except Exception as e:
            logger.error(f"Error extracting KiCad file info: {str(e)}")
            return {"error": f"Could not extract KiCad file information: {str(e)}"}

    def _get_mime_type(self, extension):
        """Get the MIME type for a file extension"""
        mime_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.pdf': 'application/pdf',
            '.zip': 'application/zip',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.md': 'text/markdown',
            '.json': 'application/json',
            '.xml': 'application/xml',
            '.html': 'text/html',
            '.css': 'text/css',
            '.js': 'application/javascript',
            '.py': 'text/x-python',
            '.c': 'text/x-c',
            '.cpp': 'text/x-c++',
            '.h': 'text/x-c',
            '.hpp': 'text/x-c++',
            '.kicad_pcb': 'application/x-kicad-pcb',
            '.kicad_sch': 'application/x-kicad-sch',
            '.kicad_pro': 'application/x-kicad-project',
            '.kicad_mod': 'application/x-kicad-footprint',
            '.kicad_sym': 'application/x-kicad-symbol',
            '.kicad_wks': 'application/x-kicad-worksheet',
            '.net': 'application/x-kicad-netlist',
            '.lib': 'application/x-kicad-library'
        }
        
        ext = extension.lower()
        return mime_map.get(ext, 'application/octet-stream')

    def on_input_key_down(self, event):
        """Handle key events in the input field"""
        keycode = event.GetKeyCode()
        if keycode == wx.WXK_RETURN:
            logger.info(f"Return key pressed, modifiers: Ctrl={event.ControlDown()}, Shift={event.ShiftDown()}")
            if event.ControlDown() or event.ShiftDown():
                # Allow newline with Ctrl+Enter or Shift+Enter
                logger.info("Adding newline with Ctrl+Enter or Shift+Enter")
                event.Skip()  # Let the TextCtrl handle the newline
            else:
                # Send message on Enter
                logger.info("Enter key pressed without modifiers, sending message")
                wx.CallAfter(self.on_send, None)  # Use CallAfter to avoid event processing issues
                return  # Don't skip the event
        else:
            # For all other keys, let the TextCtrl handle them
            event.Skip()

    def on_configure_api_key(self, event):
        """Show API key configuration dialog"""
        dlg = wx.TextEntryDialog(
            self,
            "Enter your OpenAI API key:",
            "Configure API Key",
            self.openai_manager.api_key
        )
        
        if dlg.ShowModal() == wx.ID_OK:
            api_key = dlg.GetValue().strip()
            if api_key:
                if self.openai_manager.update_api_key(api_key):
                    wx.MessageBox(
                        "API key saved successfully.",
                        "Success",
                        wx.OK | wx.ICON_INFORMATION
                    )
                else:
                    wx.MessageBox(
                        "Failed to save API key.",
                        "Error",
                        wx.OK | wx.ICON_ERROR
                    )
        
        dlg.Destroy()

    def _update_file_button_state(self):
        """Update the file button state based on whether a file is selected"""
        # This method should be called when a file is selected or cleared
        # We don't need to do anything special here since we're using a dedicated button
        pass
        
    def _show_file_drop_success(self):
        """Show a visual indicator that a file was successfully dropped"""
        # Flash the file drop area green briefly to indicate success
        original_color = self.file_drop_panel.GetBackgroundColour()
        
        def flash_green():
            self.file_drop_panel.SetBackgroundColour(wx.Colour(100, 200, 100))
            self.file_drop_panel.Refresh()
            
            # Schedule restoration of original color
            wx.CallLater(300, lambda: self.file_drop_panel.SetBackgroundColour(original_color))
            wx.CallLater(300, self.file_drop_panel.Refresh)
            
        # Use CallAfter to ensure UI update happens in the main thread
        wx.CallAfter(flash_green)

    def add_user_message_with_file(self, message, file_data):
        """Add a user message with a file attachment to the chat history"""
        # Get the current conversation
        conversation = self.conversation_manager.get_current_conversation()
        
        # Generate a unique ID for this file
        file_id = str(uuid.uuid4())
        
        # Store file data in the conversation's files dictionary
        conversation.files[file_id] = file_data
        
        # Format the message with file info
        file_name = file_data.get('name', 'unnamed file')
        formatted_message = message
        
        if not formatted_message:
            formatted_message = f"[Attached file: {file_name}]"
        else:
            formatted_message = f"{message}\n\n[Attached file: {file_name}]"
        
        # Always add file content to the message regardless of type
        # This ensures the AI can access the content
        file_content = file_data.get('content')
        if file_content and isinstance(file_content, str):
            # Limit content length for display
            content_to_display = file_content
            if len(content_to_display) > 4000:
                content_to_display = content_to_display[:4000] + "...(truncated)"
                
            formatted_message += f"\n\nFile content:\n```\n{content_to_display}\n```"
            
        # Add to the UI
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.BLACK))
        self.chat_history.AppendText("You: ")
        self.chat_history.SetDefaultStyle(wx.TextAttr(wx.BLACK))
        self.chat_history.AppendText(f"{formatted_message}\n")
        
        # Add to conversation with file reference
        message_obj = {"role": "user", "content": formatted_message, "file_ref": file_id}
        conversation.messages.append(message_obj)
        
        # Save the conversation
        self.conversation_manager.save_conversations()
        self.sidebar.update_current_conversation_title()

    def _extract_pdf_text_alternative(self, pdf_path):
        """Extract text from PDF using a basic alternative method"""
        try:
            # Try a simpler approach: just read some bytes and look for text strings
            logger.info("Attempting alternative PDF text extraction")
            
            with open(pdf_path, 'rb') as f:
                content = f.read()
                
                # Try to decode as UTF-8 with errors ignored
                text = content.decode('utf-8', errors='ignore')
                
                # Filter out non-printable characters
                import string
                printable = set(string.printable)
                text = ''.join(filter(lambda x: x in printable, text))
                
                # Clean up a bit
                import re
                # Replace multiple spaces or non-printable chars with a single space
                text = re.sub(r'\s+', ' ', text)
                
                # Extract anything that looks like reasonable text (words of 3+ chars)
                words = re.findall(r'\b[A-Za-z]{3,}\b', text)
                
                if len(words) > 10:  # If we have a reasonable amount of text
                    clean_text = ' '.join(words)
                    logger.info(f"Alternative extraction found {len(words)} words")
                    return f"Basic PDF content (extracted as text): {clean_text}"
                else:
                    logger.info("Alternative extraction didn't find meaningful text")
                    return ""
                    
        except Exception as e:
            logger.error(f"Error in alternative PDF extraction: {str(e)}")
            return ""

    def _extract_pdf_text(self, pdf_path):
        """Extract text from PDF files"""
        try:
            # Check if PyPDF2 is available
            if not self.has_pypdf2:
                logger.warning("PyPDF2 not installed. Cannot extract text from PDF.")
                return f"[PDF text extraction not available: PyPDF2 not installed]"
                
            # Try to extract text with PyPDF2
            try:
                import PyPDF2
                logger.info(f"Attempting to extract text from PDF: {pdf_path}")
                
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    page_count = len(reader.pages)
                    logger.info(f"PDF has {page_count} pages")
                    
                    # Try to extract text from each page
                    for i, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text() or ""
                            text += page_text
                            logger.info(f"Extracted {len(page_text)} characters from page {i+1}")
                        except Exception as page_error:
                            logger.error(f"Error extracting text from page {i+1}: {str(page_error)}")
                            text += f"[Error extracting text from page {i+1}]"
                    
                    logger.info(f"Total extracted text length: {len(text)}")
                    
                    # If we got empty text, provide a helpful message
                    if not text.strip():
                        logger.warning("No text extracted from PDF - may be a scanned document or image-based PDF")
                        return f"[This appears to be a PDF without extractable text. It may be a scanned document or image-based PDF.]"
                    
                    return text
            except Exception as e:
                logger.error(f"PyPDF2 extraction error: {str(e)}")
                return f"[Error extracting PDF text: {str(e)}]"
        except Exception as e:
            logger.error(f"Error in PDF extraction: {str(e)}")
            return f"[Error in PDF handling: {str(e)}]"
    
    def _get_basic_pdf_info(self, pdf_path):
        """Get basic info about a PDF file when text extraction fails"""
        try:
            file_info = []
            file_info.append(f"PDF File: {os.path.basename(pdf_path)}")
            file_info.append(f"Size: {os.path.getsize(pdf_path)/1024:.1f} KB")
            
            # Try to get more info
            with open(pdf_path, 'rb') as f:
                # Read the first few bytes to check if it's a valid PDF
                header = f.read(5)
                if header == b'%PDF-':
                    file_info.append("Valid PDF signature detected")
                    # Read a bit more to try to detect PDF version
                    f.seek(0)
                    header = f.read(10).decode('ascii', errors='ignore')
                    if 'PDF-' in header:
                        version = header.split('PDF-')[1][:3]
                        file_info.append(f"PDF Version: {version}")
                
                # Check if it might be a scanned document
                f.seek(0)
                content = f.read(10000)  # Read first 10KB
                if b'/Image' in content or b'/XObject' in content:
                    file_info.append("File appears to contain images or scanned content")
                if b'/Text' in content:
                    file_info.append("File appears to contain some text content")
                    
            return "\n".join(file_info) + "\n\nNote: This is just basic file information. Text extraction is not available because PyPDF2 is not installed."
            
        except Exception as e:
            logger.error(f"Error getting basic PDF info: {str(e)}")
            return f"PDF File: {os.path.basename(pdf_path)}\nText extraction not available (PyPDF2 not installed)"
            
    def _extract_docx_text(self, docx_path):
        """Extract text from DOCX files"""
        try:
            # Check if python-docx is available
            if not self.has_python_docx:
                logger.warning("python-docx not installed. Cannot extract text from DOCX.")
                return f"[DOCX text extraction not available: python-docx not installed]"
                
            # Try to extract text with python-docx
            try:
                import docx
                logger.info(f"Attempting to extract text from DOCX: {docx_path}")
                
                doc = docx.Document(docx_path)
                text = ""
                
                # Extract text from paragraphs
                logger.info(f"Processing {len(doc.paragraphs)} paragraphs from DOCX")
                for i, paragraph in enumerate(doc.paragraphs):
                    try:
                        if paragraph.text:
                            text += paragraph.text + "\n"
                    except Exception as para_error:
                        logger.error(f"Error extracting text from paragraph {i+1}: {str(para_error)}")
                
                # Extract text from tables
                logger.info(f"Processing {len(doc.tables)} tables from DOCX")
                for i, table in enumerate(doc.tables):
                    try:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"
                    except Exception as table_error:
                        logger.error(f"Error extracting text from table {i+1}: {str(table_error)}")
                
                logger.info(f"Total extracted text length from DOCX: {len(text)}")
                
                # If we got empty text, provide a helpful message
                if not text.strip():
                    logger.warning("No text extracted from DOCX")
                    return f"[This appears to be a DOCX without extractable text.]"
                
                return text
            except Exception as e:
                logger.error(f"python-docx extraction error: {str(e)}")
                return f"[Error extracting DOCX text: {str(e)}]"
        except Exception as e:
            logger.error(f"Error in DOCX extraction: {str(e)}")
            return f"[Error in DOCX handling: {str(e)}]"

class ChatDialog(wx.Frame):
    """Dialog window for the chat interface"""
    
    def __init__(self, parent, title):
        style = wx.DEFAULT_FRAME_STYLE | wx.RESIZE_BORDER | wx.MINIMIZE_BOX | wx.MAXIMIZE_BOX
        super(ChatDialog, self).__init__(parent, title=title, style=style)
        
        self.SetTitle("KiCad AI Plugin by ALT TAB")
        self.SetSize(700, 500)  # Wider to accommodate sidebar
        
        # Create panel and widgets
        self.chat_panel = ChatPanel(self)
        
        # Minimum size
        self.SetMinSize(wx.Size(600, 400))
        
        # Status bar
        self.CreateStatusBar()
        self.SetStatusText("Ready")
        
        # Menu
        menubar = wx.MenuBar()
        file_menu = wx.Menu()
        
        # Update help menu
        help_item = file_menu.Append(wx.ID_HELP, 'Help', 'View plugin information')
        self.Bind(wx.EVT_MENU, self.on_help, help_item)
        
        item_exit = file_menu.Append(wx.ID_EXIT, 'Exit', 'Exit application')
        self.Bind(wx.EVT_MENU, self.on_exit, item_exit)
        
        menubar.Append(file_menu, '&File')
        self.SetMenuBar(menubar)
        
        # Center on screen
        self.Centre()
        
        # Bind close event
        self.Bind(wx.EVT_CLOSE, self.on_close)
        
    def on_help(self, event):
        """Show help dialog"""
        help_text = (
            "KiCad AI Plugin by ALT TAB\n\n"
            "This plugin uses OpenAI's GPT models to assist with PCB design.\n\n"
            "Tips:\n"
            "- Use the sidebar to create new chats or switch between conversations\n"
            "- Be specific in your questions\n"
            "- You can change the AI model using the dropdown menu\n"
            "- GPT-3.5-Turbo is faster and more economical\n"
            "- GPT-4 provides more detailed responses\n\n"
            "Your chat history is saved automatically."
        )
        dlg = wx.MessageDialog(self, help_text, "About KiCad AI Plugin by ALT TAB", wx.OK | wx.ICON_INFORMATION)
        dlg.ShowModal()
        dlg.Destroy()
        
    def on_exit(self, event):
        """Handle exit from menu"""
        self.Close()
        
    def on_close(self, event):
        """Handle window close"""
        self.Hide()
        # Don't destroy, just hide
        event.Skip(False)

class AIChatPlugin(pcbnew.ActionPlugin):
    """
    KiCad plugin that adds AI capabilities
    """
    
    def __init__(self):
        super(AIChatPlugin, self).__init__()
        self.dialog = None
        self.defaults()
        
    def defaults(self):
        """Set the default plugin properties"""
        self.name = "KiCad AI Plugin by ALT TAB"
        self.category = "AI Tools"
        self.description = "AI-powered KiCad assistant by ALT TAB"
        self.show_toolbar_button = True
        
        # Set the icon path
        self.icon_file_name = os.path.join(os.path.dirname(os.path.abspath(__file__)), 
                                         "icons", "resources", "icon.png")
        
    def Run(self):
        """Run the plugin - show the chat dialog"""
        try:
            if self.dialog is None:
                logger.info("Creating chat dialog")
                parent = wx.FindWindowByName("PcbFrame")
                if not parent:
                    # Try to find the schematic frame
                    parent = wx.FindWindowByName("SchematicFrame")
                    if not parent:
                        # Fallback to None
                        parent = None
                
                self.dialog = ChatDialog(parent, "KiCad AI Plugin by ALT TAB")
                
            if self.dialog.IsShown():
                # If already visible, bring to front
                self.dialog.Raise()
            else:
                self.dialog.Show()
                
        except Exception as e:
            logger.error(f"Error running KiCad AI Plugin by ALT TAB: {str(e)}")
            wx.MessageBox(f"Error: {str(e)}", "KiCad AI Plugin by ALT TAB Error", wx.OK | wx.ICON_ERROR)

# Create the plugin instance
ai_chat_plugin = AIChatPlugin()
try:
    ai_chat_plugin.register()
except Exception:
    pass  # Fail silently if not running inside pcbnew

if __name__ == "__main__":
    # Standalone execution for testing the UI outside of KiCad
    app = wx.App(False)
    
    # Try to set up a dummy frame if needed
    frame = wx.Frame(None, title="KiCad AI Plugin Standalone Test")
    
    # Run the plugin
    plugin = AIChatPlugin()
    
    # Override dialog parent to None for standalone testing
    plugin.dialog = ChatDialog(None, "KiCad AI Plugin by ALT TAB")
    plugin.dialog.Show()
    
    app.MainLoop()